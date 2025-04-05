# ----------------------------------------------
# Title: documentModel.py
# Description: Resume and cover letter document model ***This section has not been tested***
# Author(s): Jasmine, Yui
# Date created: Mar 7, 2025
# Date modified: Mar 8, 2025
# ----------------------------------------------
import sqlite3
import pdfkit
from docx import Document
from bs4 import BeautifulSoup
import bleach
import os
import sys

class DocumentModel:
    #below are examples functions
    
    @staticmethod  
    def saveResume(userId, jobId, resumeHtml):
        pass
    #Saves the resume HTML to the database.
    
    @staticmethod
    def saveCoverLetter(userId, jobId, coverLetterHtml):
        pass
    #Saves the cover letter HTML to the database.
    
    @staticmethod
    def getResume(userId, jobId):
        pass
    #Retrieves the resume HTML from the database for a specific user and job.
    
    @staticmethod
    def getCoverLetter(userId, jobId):
        pass
    #Retrieves the cover letter HTML from the database for a specific user and job.   
    @staticmethod
    def configure_pdfkit():
        #Configure pdfkit with the path to wkhtmltopdf
        # Set the path to wkhtmltopdf based on your OS
        if sys.platform == "win32":
            # Windows path - update this to match your installation
            wkhtmltopdf_path = r'C:\Users\yuind\AppData\Local\Programs\Python\Python312\wkhtmltopdf\bin\wkhtmltopdf.exe'

        
        # Check if the file exists
        if os.path.exists(wkhtmltopdf_path):
            print(f"Using wkhtmltopdf at: {wkhtmltopdf_path}")
            # Configure pdfkit with the path
            config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
            return config
        else:
            print("wkhtmltopdf executable not found at the default location.")
            print("Please install wkhtmltopdf from: https://wkhtmltopdf.org/downloads.html")
            print("After installation, update the path in this script.")
            return None
    @staticmethod
    def generatePdf(htmlContent, outputFile):
        """
        Converts HTML content to a PDF file using pdfkit.
        """
        try:
            pdfkit.from_string(htmlContent, outputFile)
            return True
        except Exception as e:
            print(f"Error generating PDF: {e}")
            return False

    @staticmethod
    def generateWord(htmlContent, outputFile):
        """
        Converts HTML content to a Word document.
        """
        try:
            soup = BeautifulSoup(htmlContent, 'html.parser')
            doc = Document()
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'li']):
                if element.name == 'p':
                    doc.add_paragraph(element.text)
                elif element.name.startswith('h'):
                    doc.add_heading(element.text, level=int(element.name[1]))
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.text, style='ListBullet')
            doc.save(outputFile)
            return True
        except Exception as e:
            print(f"Error generating Word document: {e}")
            return False

    @staticmethod
    def sanitizeHtml(htmlContent):
        """
        Sanitizes HTML content to remove potentially harmful elements.
        """
        allowedTags = ['p', 'h1', 'h2', 'h3', 'ul', 'li', 'b', 'i', 'a']
        return bleach.clean(htmlContent, tags=allowedTags)
    
    @staticmethod
    def previewPdf(htmlContent):
        """
        Generates a PDF preview of the HTML content and returns the file path.
        """
        sanitizedHtml = DocumentModel.sanitizeHtml(htmlContent)
        outputFile = "preview.pdf"
        if DocumentModel.generatePdf(sanitizedHtml, outputFile):
            return outputFile
        return None


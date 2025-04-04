# ----------------------------------------------
# Title: main.py
# Description: Run flask with job status filtering
# Author(s): Jasmine
# Date created: Feb 28, 2025
# Date modified: Mar 10, 2025
# ----------------------------------------------
# main.py
from flask import Flask, render_template, redirect, url_for, session, request, flash
from controllers.userController import UserController
from controllers.jobController import JobController
from controllers.experienceController import ExperienceController
from models.jobModel import JobModel
from models.apiModel import apiModel
from models.documentModel import DocumentModel
from initDb import initDb, simulateUserData
from flask import request, send_file
#from weasyprint import HTML
from docx import Document
import tempfile

app = Flask(__name__)
app.secret_key = "your_secret_key"

pdfkit_config = DocumentModel.configure_pdfkit()

######### Main page ######### 
@app.route('/')
def index():
    user_id = 1  # Replace with actual session user ID
    if 'userId' not in session:
        return redirect(url_for('login'))
    user_id = session['userId']

    status_filter = request.args.get('status')

    if status_filter:
        jobs = JobModel.getJobsByStatus(user_id, status_filter)
    else:
        jobs = JobModel.getJobs(user_id)

    ai_suggestions = JobModel.getJobSuggestions(user_id)  # Fetch AI suggestions from DB

    return render_template('index.html', jobs=jobs, aiSuggestions=ai_suggestions, user_id=user_id)

@app.route('/generateSuggestions/<int:userId>')
def generate_suggestions(userId):
    # Call the function to generate new AI job suggestions
    new_suggestions = apiModel.generateJobSuggestions(userId)
    
    if new_suggestions is None:
        return jsonify({"error": "No jobs found."}), 500

    return jsonify(new_suggestions)

######### User page ######### 
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        userId = UserController.validateLogin(username, password)
        if userId:
            session['userId'] = userId
            return redirect(url_for('index'))
        else:
            flash("Invalid username or password.")
            return redirect(url_for('login'))

    return render_template('login.html')

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    if request.method == 'POST':
        return UserController.updateProfile()
    return UserController.viewProfile()

@app.route('/signUp', methods=['GET', 'POST'])
def signUp():
    if request.method == 'POST':
        full_name = request.form.get('full-name')
        email = request.form.get('email')
        phone = request.form.get('phone')
        username = request.form.get('user-name')
        password = request.form.get('password')

        return UserController.signUp(full_name, email, phone, username, password)

    return render_template('signUp.html')

@app.route('/signout')
def signout():
    return UserController.signout()

######### Job page ######### 
@app.route('/newJob', methods=['GET', 'POST'])
@app.route('/newJob/<int:jobId>', methods=['GET', 'POST'])
def newJob(jobId=None):
    if 'userId' not in session:
        return redirect(url_for('login'))
    userId = session['userId']  
    return JobController.newJob(jobId, userId)

@app.route('/deleteJob/<int:jobId>', methods=['GET'])
def deleteJob(jobId):
    if 'userId' not in session:
        return redirect(url_for('login'))
    return JobController.deleteJob(jobId)

@app.route('/archiveJob', methods=['GET', 'POST'])
def archiveJob():
    print("Job archive page")
    return render_template('archiveJob.html')



######### Expirence page ######### 
from flask import jsonify
@app.route('/newExperience', methods=['GET', 'POST'])
def newExperience():
    return ExperienceController.new_experience()

@app.route('/api/experiences', methods=['GET'])
def getExperiences():
    return ExperienceController.get_experiences()

@app.route('/deleteExperience', methods=['POST'])
def deleteExperience():
    return ExperienceController.delete_experience()

@app.route('/newWork', methods=['GET', 'POST'])
def newWork():
    if request.method == 'POST':
        return ExperienceController.addWork()
    else:
        print("Work page GET")
        return render_template('newWork.html')

@app.route('/newVolunteer', methods=['GET', 'POST'])
def newVolunteer():
    if request.method == 'POST':
        return ExperienceController.addVolunteer()
    return render_template('newVolunteer.html')

@app.route('/newProject', methods=['GET', 'POST'])
def newProject():
    if request.method == 'POST':
        return ExperienceController.addProject()
    return render_template('newProject.html')

@app.route('/newAward', methods=['GET', 'POST'])
def newAward():
    if request.method == 'POST':
        return ExperienceController.addAward()
    return render_template('newAward.html')

@app.route('/newCertification', methods=['GET', 'POST'])
def newCertification():
    if request.method == 'POST':
        return ExperienceController.addCertification()
    return render_template('newCertification.html')

@app.route('/newEducation', methods=['GET', 'POST'])
def newEducation():
    if request.method == 'POST':
        return ExperienceController.addEducation()
    return render_template('newEducation.html')

######### Documents page ######### 
@app.route('/resume', methods=['GET', 'POST'])
def resume():
    print("Resume page")
    return render_template('resume.html')

@app.route('/coverLetter', methods=['GET', 'POST'])
def coverLetter():
    print("Cover letter page")
    return render_template('coverLetter.html')

@app.route('/download', methods=['POST'])
def download():
    format_type = request.form.get('download-as')
    resume_html = request.form.get('resumeContent')
    cover_html = request.form.get('coverLetterContent')

    # Combine both sections into one HTML
    html_content = f"""
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; padding: 20px; }}
            h1 {{ color: #333; }}
            hr {{ margin: 30px 0; }}
        </style>
    </head>
    <body>
        <h1>Resume</h1>
        {resume_html}
        <hr>
        <h1>Cover Letter</h1>
        {cover_html}
    </body>
    </html>
    """

    if format_type in ['1', '3']:  # PDF or both
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        HTML(string=html_content).write_pdf(pdf_path)
        return send_file(pdf_path, as_attachment=True, download_name="documents.pdf")

    if format_type in ['2']:  # Word only
        doc = Document()
        doc.add_heading('Resume', 0)
        doc.add_paragraph(resume_html)
        doc.add_heading('Cover Letter', 0)
        doc.add_paragraph(cover_html)
        word_path = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
        doc.save(word_path)
        return send_file(word_path, as_attachment=True, download_name="documents.docx")

    return "Invalid format", 400
@app.route('/generateBoth/<int:jobId>', methods=['GET', 'POST'])
def generate_both(jobId):
    #if 'userId' not in session:
    #    return jsonify({"error": "Not logged in"}), 401
    
    userId = session['userId']
    
    try:
        # Generate resume and cover letter
        resume_html = apiModel.generateResume(userId, jobId)
        cover_letter_html = apiModel.generateCoverLetter(userId, jobId)
        
        # Return the generated content as JSON
        return jsonify({
            "resume": resume_html if resume_html else "",
            "coverLetter": cover_letter_html if cover_letter_html else ""
        })
    except Exception as e:
        print(f"Error generating documents: {str(e)}")
        return jsonify({"error": "Error generating documents"}), 500
@app.route('/testGenerate/<int:jobId>')
def test_generate(jobId):
    if 'userId' not in session:
        return "Not logged in", 401
    
    userId = session['userId']
    
    try:
        # Generate resume
        resume_html = apiModel.generateResume(userId, jobId)
        return f"Resume generation test: {'Success' if resume_html else 'Failed'}"
    except Exception as e:
        return f"Error: {str(e)}"
    
@app.route('/listJobs')
def list_jobs():
    if 'userId' not in session:
        return "Not logged in", 401
    
    userId = session['userId']
    jobs = JobModel.getJobs(userId)
    
    result = "<h1>Jobs</h1><ul>"
    for job in jobs:
        result += f"<li>ID: {job[0]} - Title: {job[2]}</li>"
    result += "</ul>"
    
    return result
if __name__ == '__main__':
    initDb()
    userId = 1
    app.run(debug=True, port=8000)

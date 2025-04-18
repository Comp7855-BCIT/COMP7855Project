# ----------------------------------------------
# Title: main.py
# Description: Run flask with job status filtering
# Author(s): Jasmine
# Date created: Feb 28, 2025
# Date modified: Mar 10, 2025
# ----------------------------------------------
# main.py

from flask import send_file
import tempfile
import os

from flask import Flask, render_template, redirect, url_for, session, request, flash, jsonify
from controllers.apiController import apiController
from controllers.documentController import DocumentController
from controllers.userController import UserController
from controllers.jobController import JobController
from controllers.experienceController import ExperienceController
# from models.jobModel import JobModel
# from models.apiModel import apiModel
# from models.documentModel import DocumentModel
from initDb import initDb, simulateUserData
from flask import request, send_file
from weasyprint import HTML
from docx import Document
import tempfile

import sqlite3

app = Flask(__name__)
app.secret_key = "your_secret_key"

pdfkit_config = DocumentController.configure_pdfkit()

######### Main page #########
@app.route('/')
def index():
    #user_id = 1  # Replace with actual session user ID
    if 'userId' not in session:
        return redirect(url_for('login'))
    user_id = session['userId']
    username = UserController.getCurrentUsername()
    status_filter = request.args.get('status')

    if status_filter:
        jobs = JobController.viewJobsByStatus(user_id, status_filter)
    else:
        jobs = JobController.viewJobs(user_id)

    ai_suggestions = JobController.viewJobSuggestions(user_id)  # Fetch AI suggestions from DB

    return render_template('index.html', jobs=jobs, aiSuggestions=ai_suggestions, user_id=user_id, username=username)

@app.route('/generateSuggestions/<int:userId>')
def generate_suggestions(userId):
    # Call the function to generate new AI job suggestions
    new_suggestions = apiController.generateJobSuggestions(userId)
    
    if new_suggestions is None:
        return jsonify({"error": "No jobs found."}), 500

    return jsonify(new_suggestions)

######### User page ######### 
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        userId = UserController.validateLogin(username, password)
        if userId:
            session['userId'] = userId
            return redirect(url_for('index'))
        else:
            flash("Invalid username or password.")
            return redirect(url_for('login'))

    return render_template('login.html')

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    if request.method == 'POST':
        return UserController.updateProfile()
    return UserController.viewProfile()

@app.route('/signUp', methods=['GET', 'POST'])
def signUp():
    if request.method == 'POST':
        full_name = request.form.get('full-name')
        email = request.form.get('email')
        phone = request.form.get('phone')
        username = request.form.get('user-name')
        password = request.form.get('password')

        return UserController.signUp(full_name, email, phone, username, password)

    return render_template('signUp.html')

@app.route('/signout')
def signout():
    return UserController.signout()

######### NEW: Delete user route #########
@app.route('/deleteUser/<int:userId>', methods=['GET'])
def deleteUserRoute(userId):
    """
    Allows the user to delete their account if authorized.
    Calls UserController.deleteAccount(...).
    """
    return UserController.deleteAccount(userId)

######### Job page #########
@app.route('/newJob', methods=['GET', 'POST'])
@app.route('/newJob/<int:jobId>', methods=['GET', 'POST'])
def newJob(jobId=None):
    if 'userId' not in session:
        return redirect(url_for('login'))
    userId = session['userId']
    
    return JobController.newJob(jobId, userId)

@app.route('/deleteJob/<int:jobId>', methods=['GET'])
def deleteJob(jobId):
    if 'userId' not in session:
        return redirect(url_for('login'))
    return JobController.deleteJob(jobId)

@app.route('/archiveJob', methods=['GET', 'POST'])
def archiveJob():
    print("Job archive page")
    return render_template('archiveJob.html')

######### Expirence page #########
from flask import jsonify
@app.route('/newExperience', methods=['GET', 'POST'])
def newExperience():
    return ExperienceController.new_experience()

@app.route('/api/experiences', methods=['GET'])
def getExperiences():
    return ExperienceController.get_experiences()

@app.route('/deleteExperience', methods=['POST'])
def deleteExperience():
    return ExperienceController.delete_experience()

@app.route('/newWork', methods=['GET', 'POST'])
def newWork():
    if request.method == 'POST':
        return ExperienceController.addWork()
    else:
        print("Work page GET")
        return render_template('newWork.html')

@app.route('/newVolunteer', methods=['GET', 'POST'])
def newVolunteer():
    if request.method == 'POST':
        return ExperienceController.addVolunteer()
    return render_template('newVolunteer.html')

@app.route('/newProject', methods=['GET', 'POST'])
def newProject():
    if request.method == 'POST':
        return ExperienceController.addProject()
    return render_template('newProject.html')

@app.route('/newAward', methods=['GET', 'POST'])
def newAward():
    if request.method == 'POST':
        return ExperienceController.addAward()
    return render_template('newAward.html')

@app.route('/newCertification', methods=['GET', 'POST'])
def newCertification():
    if request.method == 'POST':
        return ExperienceController.addCertification()
    return render_template('newCertification.html')

@app.route('/newEducation', methods=['GET', 'POST'])
def newEducation():
    if request.method == 'POST':
        return ExperienceController.addEducation()
    return render_template('newEducation.html')

@app.route('/newWorkTitle/<title>', methods=['GET', 'POST'])
def editWorkByTitle(title):
    return ExperienceController.editWorkByTitle(title)

@app.route('/newVolunteerTitle/<title>', methods=['GET', 'POST'])
def editVolunteerByTitle(title):
    return ExperienceController.editVolunteerByTitle(title)

@app.route('/newProjectTitle/<title>', methods=['GET', 'POST'])
def editProjectByTitle(title):
    return ExperienceController.editProjectByTitle(title)

@app.route('/newAwardTitle/<title>', methods=['GET', 'POST'])
def editAwardByTitle(title):
    return ExperienceController.editAwardByTitle(title)

@app.route('/newCertificationTitle/<title>', methods=['GET', 'POST'])
def editCertificationByTitle(title):
    return ExperienceController.editCertificationByTitle(title)

@app.route('/newEducationTitle/<title>', methods=['GET', 'POST'])
def editEducationByTitle(title):
    return ExperienceController.editEducationByTitle(title)

######### Documents page #########
@app.route('/download', methods=['POST'])
def download():
    format_type = request.form.get('download-as')
    resume_html = request.form.get('resumeContent')
    cover_html = request.form.get('coverLetterContent')

    # Combine both sections into one HTML
    html_content = f"""
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; padding: 20px; }}
            h1 {{ color: #333; }}
            hr {{ margin: 30px 0; }}
        </style>
    </head>
    <body>
        <h1>Resume</h1>
        {resume_html}
        <hr>
        <h1>Cover Letter</h1>
        {cover_html}
    </body>
    </html>
    """

    # NOTE: weasyprint usage is commented out, so this is unused. 

    if format_type in ['1', '3']:  # PDF or both
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        # HTML(string=html_content).write_pdf(pdf_path)
        return send_file(pdf_path, as_attachment=True, download_name="documents.pdf")

    if format_type in ['2']:  # Word only
        doc = Document()
        doc.add_heading('Resume', 0)
        doc.add_paragraph(resume_html)
        doc.add_heading('Cover Letter', 0)
        doc.add_paragraph(cover_html)
        word_path = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
        doc.save(word_path)
        return send_file(word_path, as_attachment=True, download_name="documents.docx")

    return "Invalid format", 400

@app.route('/generateBoth/<int:jobId>', methods=['GET', 'POST'])
def generate_both(jobId):
    if 'userId' not in session:
        return jsonify({"error": "Not logged in"}), 401
    
    userId = session['userId']
    
    try:
        resume_html = apiController.generateResume(userId, jobId)
        cover_letter_html = apiController.generateCoverLetter(userId, jobId)
        
        if not resume_html:
            resume_html = """
            <div style="font-family: Arial, sans-serif; padding: 20px;">
                <h1>Resume</h1>
                <p>No resume content was generated. Please try again or create manually.</p>
            </div>
            """
        
        if not cover_letter_html:
            cover_letter_html = """
            <div style="font-family: Arial, sans-serif; padding: 20px;">
                <h1>Cover Letter</h1>
                <p>No cover letter content was generated. Please try again or create manually.</p>
            </div>
            """
        
        return jsonify({
            "resume": resume_html if resume_html else "",
            "coverLetter": cover_letter_html if cover_letter_html else ""
        })
    except Exception as e:
        print(f"Error generating documents: {str(e)}")
        return jsonify({"error": "Error generating documents"}), 500
    
@app.route('/saveResume/<int:jobId>', methods=['POST'])
def save_resume(jobId):
    if 'userId' not in session:
        return jsonify({"error": "Not logged in"}), 401
    
    userId = session['userId']
    resume_content = request.json.get('resumeContent')
    cover_letter_content = request.json.get('coverLetterContent')
    
    if not resume_content:
        return jsonify({"error": "No resume content provided"}), 400
    if not cover_letter_content:
        return jsonify({"error": "No cover letter content provided"}), 400
    
    conn = None
    try:
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()

        # Save resume
        cursor.execute("SELECT id FROM resume WHERE userId = ? AND jobId = ?", (userId, jobId))
        existing_resume = cursor.fetchone()
        
        if existing_resume:
            cursor.execute("""
                UPDATE resume SET resume = ? 
                WHERE userId = ? AND jobId = ?
            """, (resume_content, userId, jobId))
        else:
            cursor.execute("""
                INSERT INTO resume (userId, jobId, resume)
                VALUES (?, ?, ?)
            """, (userId, jobId, resume_content))

        # Save cover letter
        cursor.execute("SELECT id FROM coverLetter WHERE userId = ? AND jobId = ?", (userId, jobId))
        existing_cover_letter = cursor.fetchone()

        if existing_cover_letter:
            cursor.execute("""
                UPDATE coverLetter SET coverLetter = ? 
                WHERE userId = ? AND jobId = ?
            """, (cover_letter_content, userId, jobId))
        else:
            cursor.execute("""
                INSERT INTO coverLetter (userId, jobId, coverLetter)
                VALUES (?, ?, ?)
            """, (userId, jobId, cover_letter_content))
        
        conn.commit()
        return jsonify({"success": True, "message": "Resume and Cover Letter saved successfully"})
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
    finally:
        if conn:
            conn.close()
@app.route('/downloadResume/<int:jobId>')
def download_resume(jobId):
    # Check if the user is logged in. If not, redirect to login page.
    if 'userId' not in session:
        return redirect(url_for('login'))
    
    # Retrieve the logged in user's ID from the session.
    userId = session['userId']
    
    # Generate the resume PDF file for the given user and job.
    # DocumentController.generate_pdf_from_resume is assumed to return the file path.
    pdf_path = DocumentController.generate_pdf_from_resume(userId, jobId)
    print(f"Attempting to download resume for job {jobId}")
    
    # If no PDF was generated, show an error message and redirect back to the main page.
    if not pdf_path:
        flash("Error generating resume PDF", "error")
        return redirect(url_for('index'))
    
    try:
        # Retrieve job details using the job ID.
        job = JobController.viewJobById(jobId)
        # Assuming that the job title is stored at index 2 in the job tuple,
        # assign it to 'job_title'; if not available, default to "resume".
        job_title = job[2] if job else "resume"
        
        # If job_title is empty or None, set it to "resume".
        if not job_title:
            job_title = "resume"
            
        # Create a safe filename by including only alphanumeric characters, spaces, and underscores.
        safe_filename = "".join(c for c in str(job_title) if c.isalnum() or c in (' ', '_')).rstrip()
        filename = f"{safe_filename}_resume.pdf"
        
        # Send the PDF file using Flask's send_file. It is sent as an attachment with a custom filename.
        response = send_file(pdf_path, as_attachment=True, download_name=filename, mimetype='application/pdf')
        
        # After sending the file, delete it from the file system to clean up temporary files.
        response.call_on_close(lambda: os.unlink(pdf_path))
        return response
    except Exception as e:
        print(f"Error sending file: {e}")
        flash("Error downloading resume", "error")
        return redirect(url_for('index'))
    
@app.route('/downloadCover/<int:jobId>')
def download_cover_letter(jobId):
    """
    Route to generate and download a cover letter PDF for a specific job.
    """
    # 1) Ensure the user is logged in by checking session. If not, redirect to login.
    if 'userId' not in session:
        return redirect(url_for('login'))
    
    # 2) Retrieve the user ID from session.
    userId = session['userId']
    
    # 3) Ask DocumentController to generate the PDF file path for the cover letter.
    pdf_path = DocumentController.generate_pdf_from_cover(userId, jobId)
    print(f"Attempting to download cover letter for job {jobId}")

    # 4) If PDF generation failed (pdf_path is None/empty), flash an error and redirect home.
    if not pdf_path:
        flash("Error generating cover letter PDF", "error")
        return redirect(url_for('index'))
    
    try:
        # 5) Retrieve job details for naming the file – usually job[2] is the job title.
        job = JobController.viewJobById(jobId)
        job_title = job[2] if job else "coverLetter"
        
        # 6) If we didn't get a valid title, default to "coverLetter".
        if not job_title:
            job_title = "coverLetter"
            
        # 7) Construct a safe filename by removing unwanted characters.
        safe_filename = "".join(c for c in str(job_title) if c.isalnum() or c in (' ', '_')).rstrip()
        filename = f"{safe_filename}_cover_letter.pdf"
        
        # 8) Use Flask's send_file to serve the PDF as an attachment.
        response = send_file(pdf_path, as_attachment=True, download_name=filename, mimetype='application/pdf')
        
        # 9) Once the file has been sent, delete it from disk to prevent clutter.
        response.call_on_close(lambda: os.unlink(pdf_path))
        return response
    except Exception as e:
        # 10) If anything goes wrong, log the error, flash a message, and redirect.
        print(f"Error sending file: {e}")
        flash("Error downloading cover letter", "error")
        return redirect(url_for('index'))

@app.route('/getDocuments/<int:jobId>')
def get_documents(jobId):
    """
    Route to retrieve saved resume and cover letter HTML from the database.
    Returns JSON with the 'resume' and 'coverLetter' fields.
    """
    # 1) Check if user is authenticated.
    if 'userId' not in session:
        return jsonify({"error": "Not logged in"}), 401
    
    userId = session['userId']
    try:
        # 2) Retrieve the existing resume and cover letter from DocumentController.
        resume = DocumentController.get_resume(userId, jobId)
        cover_letter = DocumentController.get_cover_letter(userId, jobId)
        
        # 3) Return them as a JSON payload; empty strings if none exist.
        return jsonify({
            "resume": resume if resume else "",
            "coverLetter": cover_letter if cover_letter else ""
        })
    except Exception as e:
        # 4) Log any error and return a JSON object with an error message.
        print(f"Error getting documents: {str(e)}")
        return jsonify({"error": "Error getting documents"}), 500  

@app.route('/testGenerate/<int:jobId>')
def test_generate(jobId):
    """
    A simple route to test AI-based resume generation for a given job.
    Returns either 'Success' or 'Failed'.
    """
    # 1) Require user to be logged in.
    if 'userId' not in session:
        return "Not logged in", 401
    
    userId = session['userId']
    try:
        # 2) Generate resume HTML via apiController.
        resume_html = apiController.generateResume(userId, jobId)
        # 3) Return 'Success' if we got HTML back, otherwise 'Failed'.
        return f"Resume generation test: {'Success' if resume_html else 'Failed'}"
    except Exception as e:
        # 4) Catch & show any errors in plain text.
        return f"Error: {str(e)}"

@app.route('/listJobs')
def list_jobs():
    """
    Route to list out all jobs for the logged-in user.
    This is mainly for debugging or quick viewing.
    """
    if 'userId' not in session:
        return "Not logged in", 401
    
    userId = session['userId']
    jobs = JobController.viewJobs(userId)
    
    # Build a simple HTML string with job info.
    result = "<h1>Jobs</h1><ul>"
    for job in jobs:
        # job is typically a tuple like (id, userId, jobTitle, ...)
        result += f"<li>ID: {job[0]} - Title: {job[2]}</li>"
    result += "</ul>"
    
    return result

if __name__ == '__main__':
    # 1) Initialize DB if needed.
    initDb()
    # userId = 1
    app.run(debug=True, port=8000)

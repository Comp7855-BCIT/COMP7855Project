# ----------------------------------------------
# Title: documentModel.py
# Description: Resume and cover letter document model
# Author(s): Jasmine, Yui
# Date created: Mar 7, 2025
# Date modified: Mar 8, 2025
# ----------------------------------------------

import sqlite3
import pdfkit
from docx import Document
from bs4 import BeautifulSoup
import bleach
import os
import sys

class DocumentModel:
    # Retrieves the resume HTML for a user and job from the database
    @staticmethod
    def get_resume(userId, jobId):
        try:
            conn = sqlite3.connect('database.db')
            cursor = conn.cursor()
            cursor.execute("""
                SELECT resume FROM resume 
                WHERE userId = ? AND jobId = ?
            """, (userId, jobId))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else None
        except Exception as e:
            print(f"Error getting resume: {e}")
            return None

    # Retrieves the cover letter HTML for a user and job from the database
    @staticmethod
    def get_cover_letter(userId, jobId):
        try:
            conn = sqlite3.connect('database.db')
            cursor = conn.cursor()
            cursor.execute("""
                SELECT coverLetter FROM coverLetter
                WHERE userId = ? AND jobId = ?
            """, (userId, jobId))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else None
        except Exception as e:
            print(f"Error getting cover letter: {e}")
            return None

    # Generates a PDF from the resume HTML stored in the database
    @staticmethod
    def generate_pdf_from_resume(userId, jobId, output_path=None):
        """
        Generates a PDF from the resume HTML.
        Returns the file path to the PDF if successful.
        """
        resume_html = DocumentModel.get_resume(userId, jobId)
        if not resume_html:
            return None

        # Create a temporary file if no output path is provided
        if not output_path:
            import tempfile
            output_path = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False).name

        config = DocumentModel.configure_pdfkit()
        if not config:
            return None

        try:
            options = {
                'encoding': 'UTF-8',
                'quiet': '',
                'margin-top': '0.5in',
                'margin-right': '0.5in',
                'margin-bottom': '0.5in',
                'margin-left': '0.5in',
            }
            pdfkit.from_string(resume_html, output_path, configuration=config, options=options)
            return output_path
        except Exception as e:
            print(f"Error generating PDF: {e}")
            return None

    # Generates a PDF from the cover letter HTML stored in the database
    @staticmethod
    def generate_pdf_from_cover(userId, jobId, output_path=None):
        """
        Generates a PDF from the cover letter HTML.
        Returns the file path to the generated PDF if successful.
        """
        cover_letter_html = DocumentModel.get_cover_letter(userId, jobId)
        if not cover_letter_html:
            return None

        # Create a temporary file if no output path is provided
        if not output_path:
            import tempfile
            output_path = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False).name

        config = DocumentModel.configure_pdfkit()
        if not config:
            return None

        try:
            options = {
                'encoding': 'UTF-8',
                'quiet': '',
                'margin-top': '0.5in',
                'margin-right': '0.5in',
                'margin-bottom': '0.5in',
                'margin-left': '0.5in',
            }
            pdfkit.from_string(cover_letter_html, output_path, configuration=config, options=options)
            return output_path
        except Exception as e:
            print(f"Error generating PDF: {e}")
            return None

    # Saves the resume HTML to the database
    @staticmethod  
    def saveResume(userId, jobId, resumeHtml):
        pass

    # Saves the cover letter HTML to the database
    @staticmethod
    def saveCoverLetter(userId, jobId, coverLetterHtml):
        pass

    # Retrieves the cover letter HTML from the database
    @staticmethod
    def getCoverLetter(userId, jobId):
        pass

    # Configures pdfkit with the wkhtmltopdf executable path
    @staticmethod
    def configure_pdfkit():
        # Set the path to wkhtmltopdf based on the operating system
        if sys.platform == "win32":
            # Windows path - update this as needed to your folder location
            wkhtmltopdf_path = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'
        else:
            # Non-Windows default path; change as needed
            wkhtmltopdf_path = '/usr/local/bin/wkhtmltopdf'

        # Check if the wkhtmltopdf executable exists at the path
        if os.path.exists(wkhtmltopdf_path):
            print(f"Using wkhtmltopdf at: {wkhtmltopdf_path}")
            config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
            return config
        else:
            print("wkhtmltopdf executable not found at the default location.")
            print("Please install wkhtmltopdf from: https://wkhtmltopdf.org/downloads.html")
            print("After installation, update the path in this script.")
            return None

    # Converts HTML content to a PDF file using pdfkit
    @staticmethod
    def generatePdf(htmlContent, outputFile):
        try:
            pdfkit.from_string(htmlContent, outputFile)
            return True
        except Exception as e:
            print(f"Error generating PDF: {e}")
            return False

    # Converts HTML content to a Word document using python-docx and BeautifulSoup
    @staticmethod
    def generateWord(htmlContent, outputFile):
        try:
            soup = BeautifulSoup(htmlContent, 'html.parser')
            doc = Document()
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'li']):
                if element.name == 'p':
                    doc.add_paragraph(element.text)
                elif element.name.startswith('h'):
                    doc.add_heading(element.text, level=int(element.name[1]))
                elif element.name == 'ul':
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.text, style='ListBullet')
            doc.save(outputFile)
            return True
        except Exception as e:
            print(f"Error generating Word document: {e}")
            return False

    # Sanitizes HTML content using bleach to remove disallowed tags
    @staticmethod
    def sanitizeHtml(htmlContent):
        allowedTags = ['p', 'h1', 'h2', 'h3', 'ul', 'li', 'b', 'i', 'a']
        return bleach.clean(htmlContent, tags=allowedTags)
    
    # Generates a PDF preview of given HTML content
    @staticmethod
    def previewPdf(htmlContent):
        sanitizedHtml = DocumentModel.sanitizeHtml(htmlContent)
        outputFile = "preview.pdf"
        if DocumentModel.generatePdf(sanitizedHtml, outputFile):
            return outputFile
        return None
